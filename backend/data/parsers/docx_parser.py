"""
DOCX Parser for extracting text from Word documents
"""

from pathlib import Path
from typing import Dict
from docx import Document
from ...utils.logger import setup_logger

logger = setup_logger("ki.parsers.docx")


class DOCXParser:
    """Parser for DOCX documents"""

    def __init__(self):
        self.supported_extensions = [".docx"]

    def parse(self, file_path: Path) -> Dict[str, any]:
        """
        Parse DOCX file and extract text

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with parsed content
        """
        try:
            logger.info(f"Parsing DOCX: {file_path.name}")

            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = []
            full_text = []

            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append({
                        "paragraph_number": i + 1,
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal"
                    })
                    full_text.append(para.text)

            # Extract tables
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                tables.append({
                    "table_number": i + 1,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0,
                    "data": table_data
                })

            combined_text = "\n\n".join(full_text)

            # Extract core properties
            metadata = self._extract_metadata(doc)

            result = {
                "file_name": file_path.name,
                "file_path": str(file_path),
                "file_type": "docx",
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables),
                "metadata": metadata,
                "paragraphs": paragraphs,
                "tables": tables,
                "full_text": combined_text,
                "char_count": len(combined_text),
                "word_count": len(combined_text.split()),
                "success": True,
                "error": None
            }

            logger.info(f"✅ Parsed {len(paragraphs)} paragraphs, {len(tables)} tables, {result['word_count']} words")

            return result

        except Exception as e:
            logger.error(f"❌ Error parsing DOCX {file_path.name}: {str(e)}")
            return {
                "file_name": file_path.name,
                "file_path": str(file_path),
                "file_type": "docx",
                "success": False,
                "error": str(e),
                "full_text": ""
            }

    def _extract_metadata(self, doc: Document) -> Dict[str, str]:
        """Extract DOCX metadata"""
        metadata = {}

        try:
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }
        except Exception as e:
            logger.warning(f"Could not extract metadata: {str(e)}")

        return metadata

    def is_supported(self, file_path: Path) -> bool:
        """Check if file is supported"""
        return file_path.suffix.lower() in self.supported_extensions


def parse_docx(file_path: Path) -> Dict[str, any]:
    """
    Convenience function to parse DOCX

    Args:
        file_path: Path to DOCX file

    Returns:
        Parsed content dictionary
    """
    parser = DOCXParser()
    return parser.parse(file_path)
